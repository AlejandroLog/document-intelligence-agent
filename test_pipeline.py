import os
import io
import docx
from minio import Minio
from dotenv import load_dotenv

from pipelines.ingest import upload_to_raw
from pipelines.preprocess import process_and_stage
from pipelines.chunking import chunk_text
from app.embeddings import create_vector_store

load_dotenv()

def main():
    test_file_path = "prueba_rag.docx"
    
    # 1. Crear documento si no existe
    if not os.path.exists(test_file_path):
        print("Creando archivo de prueba...")
        doc = docx.Document()
        doc.add_heading('Documento de Prueba RAG', 0)
        doc.add_paragraph('La Inteligencia Artificial generativa permite crear sistemas que razonan sobre documentos.')
        doc.add_paragraph('LangChain es un framework para desarrollar aplicaciones impulsadas por modelos de lenguaje.')
        doc.add_paragraph('MinIO es un sistema de almacenamiento de objetos de alto rendimiento, compatible con la API de Amazon S3.')
        doc.save(test_file_path)
        
    print("\n--- 1. Ingestión (Hacia RAW) ---")
    exito = upload_to_raw(test_file_path, test_file_path)

    if exito:
        print("\n--- 2. Procesamiento (Hacia STAGING) ---")
        process_and_stage(test_file_path)
        
        print("\n--- 3. Extracción desde STAGING ---")
        # Conectarse a MinIO para bajar el .txt
        minio_client = Minio(
            os.getenv("MINIO_ENDPOINT", "localhost:9000"),
            access_key=os.getenv("MINIO_ROOT_USER", "admin"),
            secret_key=os.getenv("MINIO_ROOT_PASSWORD", "password123"),
            secure=False
        )
        
        staging_name = "prueba_rag.txt"
        try:
            response = minio_client.get_object("staging", staging_name)
            texto_completo = response.read().decode('utf-8')
            response.close()
            response.release_conn()
            
            print(f"✅ Texto descargado de STAGING ({len(texto_completo)} caracteres).")
            
            print("\n--- 4. Chunking ---")
            # Partimos el texto en pedazos pequeños
            chunks = chunk_text(texto_completo, chunk_size=100, chunk_overlap=20)
            print(f"✅ Texto dividido en {len(chunks)} chunks.")
            
            print("\n--- 5. Vectorización y CURATED ---")
            # Generamos embeddings y guardamos en ChromaDB
            create_vector_store(chunks, document_name=test_file_path)           
            print("\n🎉 ¡Pipeline de Datos Completado con Éxito!")
            
        except Exception as e:
            print(f"❌ Error en la fase de chunking/embeddings: {e}")

if __name__ == "__main__":
    main()